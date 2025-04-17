"""
Utility functions for Azure Content Understanding
"""
import os
import json
import datetime
import time
import shutil
import re
import random
from io import BytesIO
import base64
from mimetypes import guess_type

from azure.core.credentials import AzureKeyCredential
from content_understanding.azure_content_understanding import AzureContentUnderstandingClient
from azure.search.documents.indexes import SearchIndexClient
from dotenv import load_dotenv
from langchain.schema import Document
from langchain_community.vectorstores import AzureSearch
from langchain_openai import AzureOpenAIEmbeddings
from openai import AzureOpenAI
from PIL import Image
import requests
import cv2
from moviepy.editor import VideoFileClip
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load environment variables
load_dotenv()

# Constants
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
JSON_DIR = os.path.join(SCRIPT_DIR, "json")
DOCUMENTS_DIR = os.path.join(SCRIPT_DIR, "video")
RESULTS_DIR = os.path.join(SCRIPT_DIR, "results")
IMAGES_DIR = os.path.join(SCRIPT_DIR, "frames")

# Ensure directories exist
os.makedirs(JSON_DIR, exist_ok=True)
os.makedirs(DOCUMENTS_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)

# Azure credentials
AZURE_AI_SERVICE_ENDPOINT = os.getenv("AZURE_AI_CU_ENDPOINT")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_CU_ENDPOINT")
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_SERVICE_ENDPOINT")
AZURE_SEARCH_KEY = os.getenv("AZURE_SEARCH_SERVICE_ADMIN_KEY")
AZURE_OPENAI_MODEL = os.getenv("AZURE_OPENAI_CU_MODEL")
AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_CU_EMBEDDING_DEPLOYMENT_NAME")
AZURE_SEARCH_INDEX_NAME = os.getenv("AZURE_SEARCH_CU_INDEX_NAME")
AZURE_AI_CU_API_VERSION = os.getenv("AZURE_AI__CU_API_VERSION")
AZURE_AI_CU_KEY = os.getenv("AZURE_AI_CU_KEY")

# Initialize client
def dummy_token_provider():
    return ""
def get_cu_client():
    """Initialize and return the Azure Content Understanding client."""
    return AzureContentUnderstandingClient(
        endpoint=AZURE_AI_SERVICE_ENDPOINT,
        api_version=AZURE_AI_CU_API_VERSION,
        subscription_key=AZURE_AI_CU_KEY,
        token_provider=dummy_token_provider,
        x_ms_useragent="azure-ai-content-understanding-python/search_with_video"
    )

# Initialize OpenAI client
def get_aoai_client():
    """Initialize and return the Azure OpenAI client."""
    return AzureOpenAI(
        api_key=AZURE_AI_CU_KEY,
        api_version="2024-02-15-preview",
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
    )

def convert_values_to_strings(json_obj):
    """Convert all values in a JSON object to strings."""
    return [str(value) for value in json_obj]

def remove_markdown(json_obj):
    """Remove 'markdown' keys from all segments in a JSON object."""
    for segment in json_obj:
        if 'markdown' in segment:
            del segment['markdown']
    return json_obj

def get_scene_description(scene_description):
    """Extract and format scene descriptions from video analysis results."""
    audio_visual_segments = scene_description["result"]["contents"]
    
    filtered_audio_visual_segments = remove_markdown(audio_visual_segments)
    
    audio_visual_splits = [
        "The following is a json string representing a video segment with scene description and transcript ```"
        + v + "```"
        for v in convert_values_to_strings(filtered_audio_visual_segments)
    ]
    
    docs = [Document(page_content=v) for v in audio_visual_splits]
    
    return docs

def get_fields_result(res_string):
    """Extract fields from a result string."""
    # Extract scene desc
    start_value_string = res_string.find('"valueString": "') + len('"valueString": "')
    end_value_string = res_string.find('"}', start_value_string)
    scene_desc = res_string[start_value_string:end_value_string]

    # Extract kind
    start_kind = res_string.find('"kind": "') + len('"kind": "')
    end_kind = res_string.find('"', start_kind)
    kind = res_string[start_kind:end_kind]

    # Extract startTimeMs
    start_startTimeMs = res_string.find('"startTimeMs": ') + len('"startTimeMs": ')
    end_startTimeMs = res_string.find(',', start_startTimeMs)
    startTimeMs = int(res_string[start_startTimeMs:end_startTimeMs])

    # Extract endTimeMs
    start_endTimeMs = res_string.find('"endTimeMs": ') + len('"endTimeMs": ')
    end_endTimeMs = res_string.find(',', start_endTimeMs)
    endTimeMs = int(res_string[start_endTimeMs:end_endTimeMs])

    # Extract width
    start_width = res_string.find('"width": ') + len('"width": ')
    end_width = res_string.find(',', start_width)
    width = int(res_string[start_width:end_width])

    # Extract height
    start_height = res_string.find('"height": ') + len('"height": ')
    end_height = res_string.find(',', start_height)
    height = int(res_string[start_height:end_height])

    # Extract KeyFrameTimesMs
    start_keyFrameTimesMs = res_string.find('"KeyFrameTimesMs": [') + len('"KeyFrameTimesMs": [')
    end_keyFrameTimesMs = res_string.find(']', start_keyFrameTimesMs)
    keyFrameTimesMs_str = res_string[start_keyFrameTimesMs:end_keyFrameTimesMs]
    keyFrameTimesMs = [int(x) for x in keyFrameTimesMs_str.split(',') if x.strip()]

    # Extract transcriptPhrases
    start_transcriptPhrases = res_string.find('"transcriptPhrases": [') + len('"transcriptPhrases": [')
    end_transcriptPhrases = res_string.find(']', start_transcriptPhrases)
    transcriptPhrases_str = res_string[start_transcriptPhrases:end_transcriptPhrases]
    transcriptPhrases = [x.strip() for x in transcriptPhrases_str.split(',') if x.strip()]

    return scene_desc, kind, startTimeMs, endTimeMs, width, height, keyFrameTimesMs, transcriptPhrases

def load_into_index(docs):
    """Load documents into Azure Search index using embeddings."""
    # Azure OpenAI Embeddings
    aoai_embeddings = AzureOpenAIEmbeddings(
        azure_deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME,
        openai_api_version="2024-02-15-preview",
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_key=AZURE_AI_CU_KEY
    )

    # Loading to the vector store
    vector_store = AzureSearch(
        azure_search_endpoint=AZURE_SEARCH_ENDPOINT,
        azure_search_key=AZURE_SEARCH_KEY,
        index_name=AZURE_SEARCH_INDEX_NAME,
        embedding_function=aoai_embeddings.embed_query)
    vector_store.add_documents(documents=docs)

    return vector_store

def get_index_stats(index_name):
    """Retrieve and display statistics for an Azure Search index."""
    url = AZURE_SEARCH_ENDPOINT + "/indexes/" + index_name + "/stats?api-version=2024-07-01"

    headers = {
        "Content-Type": "application/json",
        "api-key": AZURE_SEARCH_KEY,
    }
    response = requests.get(url, headers=headers)
    
    document_count = 0
    storage_size = 0
    
    if response.status_code == 200:
        res = response.json()
        document_count = res['documentCount']
        storage_size = res['storageSize']

    return document_count, storage_size

def generate_subclip(video_file, output_file, start_time_ms, end_time_ms):
    """Generate a subclip from a video file."""
    start = int(start_time_ms / 1000) - 1
    end = int(end_time_ms / 1000) + 1
    
    # Ensure start is not negative
    start = max(0, start)
    
    clip = VideoFileClip(video_file).subclip(start, end)
    clip.write_videofile(output_file, codec="libx264")

def download_frame(image_id, response):
    """Download a frame from the analysis operation."""
    cu_client = get_cu_client()
    raw_image = cu_client.get_image_from_analyze_operation(analyze_response=response, image_id=image_id)
    image = Image.open(BytesIO(raw_image))
    output_image_file = f"{IMAGES_DIR}/{image_id}.jpg"
    image.save(output_image_file, "JPEG")
    return output_image_file

def local_image_to_data_url(image_path):
    """Convert a local image to a data URL."""
    mime_type, _ = guess_type(image_path)
    if mime_type is None:
        mime_type = "application/octet-stream"
    with open(image_path, "rb") as image_file:
        base64_encoded_data = base64.b64encode(image_file.read()).decode("utf-8")
    return f"data:{mime_type};base64,{base64_encoded_data}"

def gpt4o_image(image_file, prompt):
    """Generate a description for an image using GPT-4o."""
    aoai_client = get_aoai_client()
    response = aoai_client.chat.completions.create(
        model=AZURE_OPENAI_MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant to analyze images."
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": local_image_to_data_url(image_file)}
                    }
                ]
            }
        ],
        max_tokens=800,
        temperature=0.7
    )
    return response.choices[0].message.content

def add_image_to_docx(doc_path, image_path, image_width=5):
    """Add an image to a Word document with auto-generated captions."""
    doc = DocxDocument(doc_path) if os.path.exists(doc_path) else DocxDocument()

    if image_width:
        doc.add_picture(image_path, width=Inches(image_width))
    else:
        doc.add_picture(image_path)

    # Get one line image description
    caption = gpt4o_image(image_path, "Describe this image in one line")
    doc.add_heading("Single image description", level=3)
    paragraph = doc.add_paragraph(caption)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Get detailed image description
    detailed_caption = gpt4o_image(image_path, "Describe this image in multiple lines")
    doc.add_heading("Detailed image description", level=3)
    paragraph = doc.add_paragraph(detailed_caption)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Get tags from the image
    tags = gpt4o_image(image_path, "Describe this image using some keywords and tags and emojis")
    doc.add_heading("Tags and emojis", level=3)
    doc.add_paragraph(tags)

    # Saving docx file
    doc.save(doc_path)
    
def get_jpg_files(directory, prefix):
    """Get a list of all .jpg files in a directory that start with a given prefix."""
    jpg_files_list = []
    for filename in os.listdir(directory):
        if filename.startswith(prefix) and filename.endswith('.jpg'):
            jpg_files_list.append(filename)
    return jpg_files_list

def create_video_analyzer_template():
    """Create a basic video content understanding template."""
    analyzer_template = {
        "description": "Generating content understanding from video for Real Estate.",
        "scenario": "videoShot",
        "config": {
            "returnDetails": True,
            "locales": [
                "en-US", "es-ES", "es-MX", "fr-FR", "hi-IN", "it-IT", "ja-JP",
                "ko-KR", "pt-BR", "zh-CN"
            ],
            "enableFace": False
        },
        "fieldSchema": {
            "name": "Content Understanding for Real Estate",
            "descriptions": "Generate content understanding from this video.",
            "fields": {
                "sceneDescription": {
                    "type": "string",
                    "description": "Describe the scene in detail, including any objects, people, or settings visible."
                }
            }
        }
    }
    
    template_path = os.path.join(JSON_DIR, "video_content_understanding.json")
    with open(template_path, "w") as f:
        json.dump(analyzer_template, f, indent=4)
        
    return template_path

def create_real_estate_analyzer_template():
    """Create a real estate specific content understanding template."""
    prompt = "Analyze the video and extract all relevant details about the property to create a captivating and professional real estate listing. Identify key specifications such as the number of bedrooms and bathrooms, total square footage, and location (if available). Describe the architectural style and highlight standout features, including luxury elements like high-end appliances, custom finishes, premium materials, renowned brands, and breathtaking views. Emphasize unique selling points that set the home apart, such as smart home technology, energy-efficient upgrades, or resort-style amenities. Additionally, showcase the property's lifestyle benefits, including its proximity to top-rated schools, shopping, dining, entertainment, parks, and transportation hubs. Structure the listing for maximum readability, using an engaging yet sophisticated tone that appeals to potential buyers, evoking excitement and desire for the home."
    
    analyzer_template = {
        "description": "Generating content understanding from video for Real Estate.",
        "scenario": "videoShot",
        "config": {
            "returnDetails": True,
            "locales": [
                "en-US", "es-ES", "es-MX", "fr-FR", "hi-IN", "it-IT", "ja-JP",
                "ko-KR", "pt-BR", "zh-CN"
            ],
            "enableFace": False
        },
        "fieldSchema": {
            "name": "Content Understanding for Real Estate",
            "descriptions": "Generate content understanding from this Real estate video.",
            "fields": {
                "realEstateInformation": {
                    "type": "string",
                    "description": prompt
                }
            }
        }
    }
    
    template_path = os.path.join(JSON_DIR, "real_estate.json")
    with open(template_path, "w") as f:
        json.dump(analyzer_template, f, indent=4)
        
    return template_path