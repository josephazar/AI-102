"""
Azure Content Understanding implementation for real estate video analysis
"""
import os
import time
import re
import random
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from content_understanding.content_understanding_utils import (
    get_cu_client, get_aoai_client, get_scene_description, get_fields_result, 
    load_into_index, download_frame, generate_subclip, add_image_to_docx,
    gpt4o_image, get_jpg_files, create_video_analyzer_template, create_real_estate_analyzer_template,
    JSON_DIR, DOCUMENTS_DIR, RESULTS_DIR, IMAGES_DIR, SCRIPT_DIR,
    AZURE_AI_SERVICE_ENDPOINT, AZURE_OPENAI_ENDPOINT, AZURE_SEARCH_ENDPOINT, 
    AZURE_OPENAI_MODEL, AZURE_SEARCH_INDEX_NAME, AZURE_AI_CU_API_VERSION
)

def process_video(video_file_path):
    """Process a video file to extract scene descriptions and keyframes."""
    print(f"Processing video: {video_file_path}")
    
    # Debug information
    print(f"Azure AI endpoint: {AZURE_AI_SERVICE_ENDPOINT}")
    print(f"API version: {AZURE_AI_CU_API_VERSION}")
    
    # Prepare the file
    video_file_name = os.path.basename(video_file_path)
    video_path = os.path.join(DOCUMENTS_DIR, video_file_name)
    
    # Copy the file to our documents directory if it's not already there
    if video_file_path != video_path and not os.path.exists(video_path):
        import shutil
        shutil.copy(video_file_path, video_path)
    
    # Create unique analyzer ID
    analyzer_id = f"videoanalyzer_{time.strftime('%Y%m%d%H%M%S')}"
    
    # Print the URL that will be used
    analyzer_url = f"{AZURE_AI_SERVICE_ENDPOINT}/contentunderstanding/analyzers/{analyzer_id}?api-version={AZURE_AI_CU_API_VERSION}"
    print(f"Content Understanding analyzer URL: {analyzer_url}")
    
    # Get client
    cu_client = get_cu_client()
    
    # Create analyzer template
    analyzer_template_path = create_video_analyzer_template()
    
    print("Creating analyzer...")
    create_response = cu_client.begin_create_analyzer(analyzer_id, analyzer_template_path=analyzer_template_path)
    result = cu_client.poll_result(create_response)
    
    print("Analyzing video...")
    start = time.time()
    
    # Submit the video for content analysis
    response = cu_client.begin_analyze(analyzer_id, file_location=video_path)
    
    # Poll until complete
    video_result = cu_client.poll_result(response, timeout_seconds=3600)
    
    elapsed = time.time() - start
    print(f"Analysis completed in {time.strftime('%H:%M:%S', time.gmtime(elapsed))}")
    
    # Process the scene descriptions
    segments = get_scene_description(video_result)
    print(f"Extracted {len(segments)} scene segments")
    
    # Index the segments
    print("Indexing segments in Azure Search...")
    vector_store = load_into_index(segments)
    
    # Get and download keyframes
    keyframe_ids = set()
    result_data = video_result.get("result", {})
    contents = result_data.get("contents", [])
    
    for content in contents:
        markdown_content = content.get("markdown", "")
        if isinstance(markdown_content, str):
            keyframe_ids.update(re.findall(r"(keyFrame\.\d+)", markdown_content))
    
    print(f"Downloading {len(keyframe_ids)} keyframes...")
    keyframe_files = []
    for keyframe_id in keyframe_ids:
        output_file = download_frame(keyframe_id, response)
        keyframe_files.append(output_file)
    
    # Delete the analyzer when done
    cu_client.delete_analyzer(analyzer_id)
    
    return vector_store, video_result, keyframe_files

def generate_real_estate_listing(video_file_path):
    """Generate a real estate listing from a video file."""
    print(f"Generating real estate listing for video: {video_file_path}")
    
    # Prepare the file
    video_file_name = os.path.basename(video_file_path)
    video_path = os.path.join(DOCUMENTS_DIR, video_file_name)
    
    # Copy the file to our documents directory if it's not already there
    if video_file_path != video_path and not os.path.exists(video_path):
        import shutil
        shutil.copy(video_file_path, video_path)
    
    # Create unique analyzer ID
    analyzer_id = f"realestate_{time.strftime('%Y%m%d%H%M%S')}"
    
    # Get clients
    cu_client = get_cu_client()
    aoai_client = get_aoai_client()
    
    # Create analyzer template
    analyzer_template_path = create_real_estate_analyzer_template()
    
    print("Creating real estate analyzer...")
    create_response = cu_client.begin_create_analyzer(analyzer_id, analyzer_template_path=analyzer_template_path)
    result = cu_client.poll_result(create_response)
    
    print("Analyzing video for real estate listing...")
    start = time.time()
    
    # Submit the video for content analysis
    response = cu_client.begin_analyze(analyzer_id, file_location=video_path)
    
    # Poll until complete
    video_result = cu_client.poll_result(response, timeout_seconds=3600)
    
    elapsed = time.time() - start
    print(f"Analysis completed in {time.strftime('%H:%M:%S', time.gmtime(elapsed))}")
    
    # Parse the results
    listing_text = ""
    if "result" in video_result and "contents" in video_result["result"]:
        for content in video_result["result"]["contents"]:
            if "fields" in content and "realEstateInformation" in content["fields"]:
                listing_text += content["fields"]["realEstateInformation"]["valueString"] + " "
    
    # Generate final listing with OpenAI
    print("Generating final listing with OpenAI...")
    prompt = "You are an AI assistant for real estate companies. Create a comprehensive real estate listing based on the following video analysis."
    
    result = aoai_client.chat.completions.create(
        model=AZURE_OPENAI_MODEL,
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": listing_text}
        ],
        max_tokens=1500,
        temperature=0.7
    )
    
    listing_content = result.choices[0].message.content
    
    # Get and download keyframes
    keyframe_ids = set()
    result_data = video_result.get("result", {})
    contents = result_data.get("contents", [])
    
    for content in contents:
        markdown_content = content.get("markdown", "")
        if isinstance(markdown_content, str):
            keyframe_ids.update(re.findall(r"(keyFrame\.\d+)", markdown_content))
    
    print(f"Downloading {len(keyframe_ids)} keyframes...")
    keyframe_files = []
    for keyframe_id in keyframe_ids:
        output_file = download_frame(keyframe_id, response)
        keyframe_files.append(output_file)
    
    # Create the DOCX file
    doc_name = f"real_estate_listing_{os.path.splitext(video_file_name)[0]}.docx"
    doc_path = os.path.join(RESULTS_DIR, doc_name)
    
    doc = DocxDocument()
    
    # Add header
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.text = "Real estate listing document"
    
    # Add content
    doc.add_heading("Real estate listing", level=1)
    doc.add_heading("Description", level=2)
    doc.add_paragraph(listing_content)
    doc.add_heading("Images", level=2)
    
    # Save the initial document
    doc.save(doc_path)
    
    # Pick random images to include (max 5 for time constraints)
    if keyframe_files:
        selected_images = random.sample(keyframe_files, min(5, len(keyframe_files)))
        
        # Add the images with captions
        for img_file in selected_images:
            add_image_to_docx(doc_path, img_file, image_width=6)
    
    # Delete the analyzer when done
    cu_client.delete_analyzer(analyzer_id)
    
    print(f"Listing generated and saved to {doc_path}")
    return doc_path, keyframe_files

def search_video_content(query, vector_store, top_k=3):
    """Search for video segments matching a query."""
    results = vector_store.similarity_search(query=query, k=top_k)
    
    parsed_results = []
    for doc in results:
        res_string = doc.page_content.split("```")[1].replace("'", "\"")
        scene_desc, kind, startTimeMs, endTimeMs, width, height, keyFrameTimesMs, transcriptPhrases = get_fields_result(res_string)
        
        parsed_results.append({
            "scene_description": scene_desc,
            "kind": kind,
            "start_time_ms": startTimeMs,
            "end_time_ms": endTimeMs,
            "width": width,
            "height": height,
            "key_frame_times_ms": keyFrameTimesMs,
            "transcript_phrases": transcriptPhrases
        })
    
    return parsed_results

def chat_with_video(query, vector_store):
    """ChatGPT-enhanced search of video content."""
    # First, search the vector store for relevant segments
    search_results = search_video_content(query, vector_store)
    
    if not search_results:
        return "I couldn't find any relevant information in the video for your question."
    
    # Get client
    aoai_client = get_aoai_client()
    
    # Format the results for OpenAI
    context = "Video segments:\n\n"
    for i, result in enumerate(search_results):
        context += f"Segment {i+1}:\n"
        context += f"Description: {result['scene_description']}\n"
        if result['transcript_phrases']:
            context += f"Transcript: {', '.join(result['transcript_phrases'])}\n"
        context += "\n"
    
    # Ask OpenAI to answer based on the retrieved content
    response = aoai_client.chat.completions.create(
        model=AZURE_OPENAI_MODEL,
        messages=[
            {
                "role": "system", 
                "content": "You are an assistant that answers questions about a video based on extracted scenes. Use only the information provided in the context. If you don't know the answer, say so."
            },
            {
                "role": "user",
                "content": f"Question: {query}\n\nContext:\n{context}"
            }
        ],
        max_tokens=800,
        temperature=0.5
    )
    
    return response.choices[0].message.content

def generate_summary(video_result):
    """Generate a summary of the video content using OpenAI."""
    # Extract all scene descriptions
    descriptions = []
    if "result" in video_result and "contents" in video_result["result"]:
        for content in video_result["result"]["contents"]:
            if "fields" in content and "sceneDescription" in content["fields"]:
                descriptions.append(content["fields"]["sceneDescription"]["valueString"])
    
    context = "Video scene descriptions:\n\n" + "\n".join(descriptions)
    
    # Get client
    aoai_client = get_aoai_client()
    
    # Get OpenAI to summarize
    response = aoai_client.chat.completions.create(
        model=AZURE_OPENAI_MODEL,
        messages=[
            {
                "role": "system", 
                "content": "You are an assistant that summarizes real estate videos. Create a concise summary that highlights the key features and areas shown in the video."
            },
            {
                "role": "user",
                "content": f"Please summarize this video based on the following scene descriptions:\n\n{context}"
            }
        ],
        max_tokens=500,
        temperature=0.5
    )
    
    return response.choices[0].message.content

def generate_report(video_result, keyframe_files, summary):
    """Generate a comprehensive report of the video analysis."""
    doc_name = f"video_analysis_report_{time.strftime('%Y%m%d%H%M%S')}.docx"
    doc_path = os.path.join(RESULTS_DIR, doc_name)
    
    doc = DocxDocument()
    
    # Add heading
    doc.add_heading("Video Analysis Report", level=1)
    
    # Add summary
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary)
    
    # Add scene breakdown
    doc.add_heading("Scene Breakdown", level=2)
    if "result" in video_result and "contents" in video_result["result"]:
        for i, content in enumerate(video_result["result"]["contents"]):
            if "fields" in content and "sceneDescription" in content["fields"]:
                doc.add_heading(f"Scene {i+1}", level=3)
                desc = content["fields"]["sceneDescription"]["valueString"]
                doc.add_paragraph(desc)
    
    # Add images
    doc.add_heading("Key Frames", level=2)
    if keyframe_files:
        selected_images = random.sample(keyframe_files, min(5, len(keyframe_files)))
        for img_file in selected_images:
            doc.add_picture(img_file, width=Inches(6))
            caption = gpt4o_image(img_file, "Describe this image in one line")
            paragraph = doc.add_paragraph(caption)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Save document
    doc.save(doc_path)
    return doc_path

if __name__ == "__main__":
    # Example usage
    video_file = os.path.join(DOCUMENTS_DIR, "paris.mp4")
    
    # Check if the file exists
    print(f"Looking for video file at: {video_file}")
    print(f"File exists: {os.path.exists(video_file)}")
    
    if os.path.exists(video_file):
        print("Processing video for Azure Content Understanding...")
        vector_store, video_result, keyframe_files = process_video(video_file)
        
        print("\nGenerating real estate listing...")
        doc_path, _ = generate_real_estate_listing(video_file)
        
        print("\nTesting video search...")
        results = search_video_content("living room", vector_store)
        for i, result in enumerate(results):
            print(f"Result {i+1}:")
            print(f"Description: {result['scene_description']}")
            print(f"Start time: {result['start_time_ms']} ms")
            print()
        
        print("\nDone! Check the results directory for output files.")
    else:
        print(f"Error: Video file {video_file} not found!")