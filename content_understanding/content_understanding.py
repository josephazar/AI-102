import os
import json
import datetime
import shutil
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from azure_content_understanding import AzureContentUnderstandingClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexClient
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from langchain.schema import Document
from langchain.vectorstores.azuresearch import AzureSearch
from langchain_openai import AzureOpenAIEmbeddings
from openai import AzureOpenAI
from PIL import Image
import requests
from io import BytesIO
import base64
import cv2
from moviepy.editor import VideoFileClip
import re
import random

# Load environment variables
load_dotenv()

# Constants
JSON_DIR = "json"
DOCUMENTS_DIR = "video"
RESULTS_DIR = "results"
IMAGES_DIR = "frames"

# Ensure directories exist
os.makedirs(RESULTS_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)

# Azure credentials
AZURE_AI_SERVICE_ENDPOINT = os.getenv("AZURE_AI_CU_ENDPOINT")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_CU_ENDPOINT")
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_SERVICE_ENDPOINT")
AZURE_SEARCH_KEY = os.getenv("AZURE_SEARCH_SERVICE_ADMIN_KEY")
AZURE_OPENAI_MODEL = os.getenv("AZURE_OPENAI_CU_MODEL")
AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_CU_EMBEDDING_DEPLOYMENT_NAME")
AZURE_SEARCH_INDEX_NAME = os.getenv("AZURE_SEARCH_CU_INDEX_NAME")
AZURE_AI_CU_API_VERSION = os.getenv("AZURE_AI__CU_API_VERSION")

# Initialize clients
credential = DefaultAzureCredential()
token_provider = get_bearer_token_provider(credential, "https://cognitiveservices.azure.com/.default")

cu_client = AzureContentUnderstandingClient(
    endpoint=AZURE_AI_SERVICE_ENDPOINT,
    api_version=AZURE_AI_CU_API_VERSION,
    token_provider=token_provider,
    x_ms_useragent="azure-ai-content-understanding-python/search_with_video"
)

search_client = SearchIndexClient(endpoint=AZURE_SEARCH_ENDPOINT, credential=AzureKeyCredential(AZURE_SEARCH_KEY))

aoai_client = AzureOpenAI(
    api_key=os.getenv("AZURE_AI_CU_KEY"),
    api_version="2024-10-21",
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
)

def download_frame(url, output_path):
    response = requests.get(url)
    img = Image.open(BytesIO(response.content))
    img.save(output_path)
    return output_path

def get_fields_result(res_string):
    res_json = json.loads(res_string)
    scene_desc = res_json["sceneDescription"]
    kind = res_json["kind"]
    startTimeMs = res_json["startTimeMs"]
    endTimeMs = res_json["endTimeMs"]
    width = res_json["width"]
    height = res_json["height"]
    keyFrameTimesMs = res_json["keyFrameTimesMs"]
    transcriptPhrases = res_json["transcriptPhrases"]
    return scene_desc, kind, startTimeMs, endTimeMs, width, height, keyFrameTimesMs, transcriptPhrases

def load_into_index(docs):
    aoai_embeddings = AzureOpenAIEmbeddings(
        azure_deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME,
        openai_api_version="2025-01-01-preview",
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        azure_ad_token_provider=token_provider
    )
    vector_store = AzureSearch(
        azure_search_endpoint=AZURE_SEARCH_ENDPOINT,
        azure_search_key=AZURE_SEARCH_KEY,
        index_name=AZURE_SEARCH_INDEX_NAME,
        embedding_function=aoai_embeddings.embed_query
    )
    vector_store.add_documents(documents=docs)
    return vector_store

def process_video(video_file_path):
    video_file_name = os.path.splitext(os.path.basename(video_file_path))[0]
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    job_name = f"{video_file_name}_{timestamp}"
    
    shutil.copy(video_file_path, os.path.join(DOCUMENTS_DIR, os.path.basename(video_file_path)))
    video_doc_path = os.path.join(DOCUMENTS_DIR, os.path.basename(video_file_path))

    ANALYZER_TEMPLATE_PATH = os.path.join(JSON_DIR, "video_content_understanding.json")
    with open(ANALYZER_TEMPLATE_PATH, "r") as f:
        analyzer_template = json.load(f)

    analyzer = cu_client.begin_analyze(
        job_name=job_name,
        video_path=video_doc_path,
        analyzer_template=analyzer_template
    ).result()

    video_result = analyzer.as_dict()
    segments = []
    keyframe_files = []
    for i, scene in enumerate(video_result["results"]["scenes"]):
        res_string = json.dumps(scene, ensure_ascii=False)
        scene_desc, kind, startTimeMs, endTimeMs, width, height, keyFrameTimesMs, transcriptPhrases = get_fields_result(res_string)
        doc = Document(page_content=f"Scene {i+1}: ```{res_string}```", metadata={"scene_id": i+1})
        segments.append(doc)
        
        for j, kf_url in enumerate(scene["keyFrameThumbnailUrls"]):
            kf_path = os.path.join(IMAGES_DIR, f"{video_file_name}_scene{i+1}_kf{j+1}.jpg")
            keyframe_files.append(download_frame(kf_url, kf_path))

    vector_store = load_into_index(segments)
    return vector_store, video_result, keyframe_files

def generate_listing(vector_store, video_result):
    prompt = """
    You are a real estate agent tasked with creating a detailed property listing based on video analysis. Using the provided video analysis results, generate a comprehensive real estate listing for a property in Paris. Include the following sections: Title, Property Overview, Key Features, Detailed Description, and Contact Information. Ensure the listing is professional, engaging, and highlights unique aspects derived from the video analysis. Format the output in plain text suitable for a Word document.
    """
    scenes = video_result["results"]["scenes"]
    scene_descriptions = [json.dumps(scene, ensure_ascii=False) for scene in scenes]
    response = aoai_client.chat.completions.create(
        model=AZURE_OPENAI_MODEL,
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": "\n".join(scene_descriptions)}
        ]
    )
    listing_text = response.choices[0].message.content

    doc = Document()
    doc.add_heading("Real Estate Listing", 0)
    sections = listing_text.split("\n\n")
    for section in sections:
        if section.strip():
            if section.startswith("Title:"):
                doc.add_heading(section.replace("Title:", "").strip(), level=1)
            elif section.startswith(("Property Overview:", "Key Features:", "Detailed Description:", "Contact Information:")):
                doc.add_heading(section.split(":")[0], level=2)
                doc.add_paragraph(section.split(":", 1)[1].strip())
            else:
                doc.add_paragraph(section.strip())

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_file = os.path.join(RESULTS_DIR, f"real_estate_listing_{timestamp}.docx")
    doc.save(docx_file)
    return docx_file

def process_and_generate(video_file):
    vector_store, video_result, keyframe_files = process_video(video_file)
    docx_file = generate_listing(vector_store, video_result)
    return docx_file, keyframe_files

def gpt4o_image(image_path, prompt):
    with open(image_path, "rb") as image_file:
        image_data = base64.b64encode(image_file.read()).decode("utf-8")
    response = aoai_client.chat.completions.create(
        model=AZURE_OPENAI_MODEL,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
            ]}
        ],
        max_tokens=500
    )
    return response.choices[0].message.content

if __name__ == "__main__":
    video_file = "video/paris.mp4"
    docx_file, keyframe_files = process_and_generate(video_file)
    print(f"Generated listing: {docx_file}")
    print(f"Keyframes saved: {keyframe_files}")