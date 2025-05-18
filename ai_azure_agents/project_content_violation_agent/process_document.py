import os
import sys
import argparse
from typing import Dict, List, Optional, Any, Tuple
import json
from datetime import datetime
import time

from dotenv import load_dotenv
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
from azure.ai.projects import AIProjectClient
from azure.identity import DefaultAzureCredential
from azure.ai.projects.models import MessageAttachment, FileSearchTool, FilePurpose

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

from config_manager import load_config, update_config

# Load environment variables
load_dotenv()

def get_document_intelligence_client() -> Optional[DocumentIntelligenceClient]:
    """
    Initialize and return a Document Intelligence client
    
    Returns:
        DocumentIntelligenceClient or None: Initialized client or None if credentials not found
    """
    endpoint = os.getenv("DOCUMENT_INTELLIGENCE_ENDPOINT")
    key = os.getenv("DOCUMENT_INTELLIGENCE_KEY")
    
    if not endpoint or not key:
        print("Error: Document Intelligence credentials not found in environment variables!")
        return None
    
    try:
        credential = AzureKeyCredential(key)
        client = DocumentIntelligenceClient(
            endpoint=endpoint,
            credential=credential
        )
        return client
    except Exception as e:
        print(f"Error initializing Document Intelligence client: {str(e)}")
        return None

def extract_text_from_pdf(document_path: str) -> List[str]:
    """
    Extract text from PDF using Azure Document Intelligence
    
    Args:
        document_path: Path to the PDF file
        
    Returns:
        List of strings where each string contains text from a page
    """
    client = get_document_intelligence_client()
    if not client:
        sys.exit(1)
    
    if not os.path.isfile(document_path):
        print(f"Error: Document not found at {document_path}")
        sys.exit(1)
    
    try:
        print(f"Processing document: {document_path}")
        with open(document_path, "rb") as document:
            poller = client.begin_analyze_document("prebuilt-read", document)
        
        result = poller.result()
        
        # Extract text by page
        pages_text = []
        for page_num, page in enumerate(result.pages, 1):
            page_text = f"******* PAGE {page_num} *******\n"
            
            # Extract text from lines
            for line in page.lines:
                page_text += f"{line.content}\n"
            
            pages_text.append(page_text)
        
        return pages_text
    
    except Exception as e:
        print(f"Error extracting text from document: {str(e)}")
        sys.exit(1)

def get_ai_project_client() -> AIProjectClient:
    """
    Initialize and return an AI Project client
    
    Returns:
        AIProjectClient: Initialized client
    """
    connection_string = os.getenv("AIPROJECT_CONNECTION_STRING")
    if not connection_string:
        print("Error: PROJECT_CONNECTION_STRING environment variable not set.")
        sys.exit(1)
    
    try:
        return AIProjectClient.from_connection_string(
            credential=DefaultAzureCredential(),
            conn_str=connection_string,
        )
    except Exception as e:
        print(f"Error initializing AI Project client: {str(e)}")
        sys.exit(1)

def check_for_violations(
    project_client: AIProjectClient,
    agent_id: str,
    thread_id: str,
    text_batch: str,
    rule_number: int,
    rule_name: str
) -> List[Dict[str, Any]]:
    """
    Check for violations of a specific rule in a batch of text
    
    Args:
        project_client: AI Project client
        agent_id: ID of the agent to use
        thread_id: ID of the thread to use
        text_batch: Batch of text to check
        rule_number: Rule number to check
        rule_name: Rule name to check
    
    Returns:
        List of violations found
    """
    prompt = f"""
    Analyze the following text for violations of Rule {rule_number}: {rule_name}.
    
    For each violation found, provide the following details in JSON format:
    - page_number: the page number (extract from the '******* PAGE X *******' markers)
    - passage: the exact text that violates the rule
    - reason: detailed explanation of why this passage violates the rule
    - confidence: either "HIGH" (clear violation) or "LOW" (potential violation requiring manual verification)
    
    Return ONLY a valid JSON array of violations. If no violations are found, return an empty array [].
    
    Here's the text to analyze:
    
    {text_batch}
    """
    
    # Create a message
    message = project_client.agents.create_message(
        thread_id=thread_id,
        role="user",
        content=prompt
    )
    
    # Process the message with the agent
    run = project_client.agents.create_and_process_run(
        thread_id=thread_id,
        agent_id=agent_id
    )
    
    # Get the response
    messages = project_client.agents.list_messages(thread_id=thread_id)
    response_text = messages.data[0].content[0].text.value
    
    # Extract JSON from response
    try:
        # Look for JSON array in the response
        json_start = response_text.find('[')
        json_end = response_text.rfind(']') + 1
        
        if json_start >= 0 and json_end > json_start:
            json_str = response_text[json_start:json_end]
            violations = json.loads(json_str)
        else:
            # If no JSON array found, try to parse the entire response
            violations = json.loads(response_text)
        
        # Ensure violations is a list
        if not isinstance(violations, list):
            violations = []
            
        # Add rule information to each violation
        for violation in violations:
            violation["rule_number"] = rule_number
            violation["rule_name"] = rule_name
        
        return violations
    
    except json.JSONDecodeError:
        print(f"Warning: Failed to parse JSON response for Rule {rule_number}.")
        print("Response:", response_text)
        return []

def create_new_thread(project_client: AIProjectClient) -> str:
    """
    Create a new thread for the agent
    
    Args:
        project_client: AI Project client
    
    Returns:
        Thread ID
    """
    thread = project_client.agents.create_thread()
    thread_id = thread.id
    print(f"Created new thread, thread ID: {thread_id}")
    
    # Update the thread ID in the config
    update_config("thread_id", thread_id)
    
    return thread_id

def delete_thread(project_client: AIProjectClient, thread_id: str) -> None:
    """
    Delete a thread
    
    Args:
        project_client: AI Project client
        thread_id: ID of the thread to delete
    """
    try:
        project_client.agents.delete_thread(thread_id)
        print(f"Deleted thread with ID: {thread_id}")
    except Exception as e:
        print(f"Error deleting thread: {str(e)}")

def process_document(document_path: str, output_format: str = "docx", delete_threads: bool = True) -> str:

    """
    Process a document for content code violations
    
    Args:
        document_path: Path to the document to process
        output_format: Format of the output report (docx or pdf)
    
    Returns:
        Path to the generated report
    """
    # Load configuration
    config = load_config()
    agent_id = config.get("agent_id")
    
    if not agent_id:
        print("Error: Agent not created or configured properly.")
        sys.exit(1)
    
    # Initialize AI Project client
    project_client = get_ai_project_client()
    
    # Create a new thread for this document processing
    thread_id = create_new_thread(project_client)
    
    # Extract text from document
    pages = extract_text_from_pdf(document_path)
    total_pages = len(pages)
    print(f"Extracted text from {total_pages} pages.")
    
    # Define rule names
    rules = {
        1: "The Protection of Children",
        2: "Harm and Offense",
        3: "Crime",
        4: "Religion",
        5: "Accuracy and Impartiality",
        6: "Fairness",
        7: "Privacy",
        8: "Interactivity",
        9: "Arrangements for Funding Content",
        10: "Advertising"
    }
    
    # Process document in batches of 10 pages
    all_violations = []
    batch_size = 10
    
    for batch_start in range(0, total_pages, batch_size):
        batch_end = min(batch_start + batch_size, total_pages)
        batch_pages = pages[batch_start:batch_end]
        batch_text = "\n\n".join(batch_pages)
        
        print(f"\nProcessing pages {batch_start + 1} to {batch_end}...")
        
        # Check for violations of each rule
        for rule_number, rule_name in rules.items():
            print(f"  Checking for violations of Rule {rule_number}: {rule_name}...")
            
            violations = check_for_violations(
                project_client,
                agent_id,
                thread_id,
                batch_text,
                rule_number,
                rule_name
            )
            
            if violations:
                print(f"    Found {len(violations)} violations.")
                all_violations.extend(violations)
            else:
                print(f"    No violations found.")
                
        # Delete the thread after processing this batch if requested
        if delete_threads:
            delete_thread(project_client, thread_id)
            # Create a new thread for the next batch
            if batch_end < total_pages:
                thread_id = create_new_thread(project_client)
    
    # Group violations by rule
    violations_by_rule = {}
    for rule_number in range(1, 11):
        violations_by_rule[rule_number] = []
    
    for violation in all_violations:
        rule_number = violation.get("rule_number")
        if rule_number in violations_by_rule:
            violations_by_rule[rule_number].append(violation)
    
    # Generate report
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = os.path.splitext(os.path.basename(document_path))[0]
    output_filename = f"{base_filename}_violations_{timestamp}"
    
    if output_format.lower() == "pdf":
        output_path = generate_pdf_report(output_filename, violations_by_rule, rules)
    else:
        output_path = generate_docx_report(output_filename, violations_by_rule, rules)
    
    print(f"\nReport generated: {output_path}")
    return output_path

def generate_docx_report(filename: str, violations_by_rule: Dict[int, List[Dict[str, Any]]], rules: Dict[int, str]) -> str:
    """
    Generate a Word document report of violations
    
    Args:
        filename: Base filename for the report
        violations_by_rule: Dictionary of violations grouped by rule
        rules: Dictionary of rule names
    
    Returns:
        Path to the generated report
    """
    # Create output directory if it doesn't exist
    output_dir = "reports"
    os.makedirs(output_dir, exist_ok=True)
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Content Code Violations Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Add summary
    total_violations = sum(len(violations) for violations in violations_by_rule.values())
    summary = doc.add_paragraph()
    summary.add_run(f"Total violations found: {total_violations}").bold = True
    doc.add_paragraph()
    
    # Add violations by rule
    for rule_number, violations in violations_by_rule.items():
        if violations:
            # Add rule heading
            rule_name = rules.get(rule_number, f"Rule {rule_number}")
            doc.add_heading(f"Rule {rule_number}: {rule_name}", 1)
            
            # Add violation count
            doc.add_paragraph(f"Violations found: {len(violations)}")
            
            # Add each violation
            for i, violation in enumerate(violations, 1):
                doc.add_heading(f"Violation {i}", 2)
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Set column widths
                table.columns[0].width = Inches(1.5)
                table.columns[1].width = Inches(5.0)
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Field"
                header_cells[1].text = "Content"
                
                # Add data rows
                for field in ["page_number", "passage", "reason", "confidence"]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = field.replace('_', ' ').title()
                    row_cells[1].text = str(violation.get(field, ""))
                
                doc.add_paragraph()
            
            doc.add_paragraph()
    
    # Save the document
    output_path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(output_path)
    
    return output_path

def generate_pdf_report(filename: str, violations_by_rule: Dict[int, List[Dict[str, Any]]], rules: Dict[int, str]) -> str:
    """
    Generate a PDF report of violations
    
    Args:
        filename: Base filename for the report
        violations_by_rule: Dictionary of violations grouped by rule
        rules: Dictionary of rule names
    
    Returns:
        Path to the generated report
    """
    # Create output directory if it doesn't exist
    output_dir = "reports"
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = os.path.join(output_dir, f"{filename}.pdf")
    
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        alignment=1,  # 0=left, 1=center, 2=right
        spaceAfter=12
    )
    
    heading1_style = ParagraphStyle(
        'Heading1',
        parent=styles['Heading1'],
        spaceAfter=10
    )
    
    heading2_style = ParagraphStyle(
        'Heading2',
        parent=styles['Heading2'],
        spaceAfter=6
    )
    
    # Create report content
    content = []
    
    # Add title
    content.append(Paragraph('Content Code Violations Report', title_style))
    content.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    content.append(Spacer(1, 12))
    
    # Add summary
    total_violations = sum(len(violations) for violations in violations_by_rule.values())
    content.append(Paragraph(f"Total violations found: {total_violations}", styles['Bold']))
    content.append(Spacer(1, 12))
    
    # Add violations by rule
    for rule_number, violations in violations_by_rule.items():
        if violations:
            # Add rule heading
            rule_name = rules.get(rule_number, f"Rule {rule_number}")
            content.append(Paragraph(f"Rule {rule_number}: {rule_name}", heading1_style))
            
            # Add violation count
            content.append(Paragraph(f"Violations found: {len(violations)}", styles['Normal']))
            content.append(Spacer(1, 6))
            
            # Add each violation
            for i, violation in enumerate(violations, 1):
                content.append(Paragraph(f"Violation {i}", heading2_style))
                
                # Create table data
                table_data = [
                    ["Field", "Content"],
                    ["Page Number", str(violation.get("page_number", ""))],
                    ["Passage", str(violation.get("passage", ""))],
                    ["Reason", str(violation.get("reason", ""))],
                    ["Confidence", str(violation.get("confidence", ""))]
                ]
                
                # Create table
                table = Table(table_data, colWidths=[100, 400])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 12),
                    ('BACKGROUND', (0, 1), (1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                content.append(table)
                content.append(Spacer(1, 12))
    
    # Build the PDF
    doc.build(content)
    
    return output_path

def main():
    parser = argparse.ArgumentParser(description="Process a document for content code violations")
    parser.add_argument("document_path", help="Path to the document to process")
    parser.add_argument(
        "--format", "-f",
        choices=["docx", "pdf"],
        default="docx",
        help="Format of the output report (docx or pdf)"
    )
    parser.add_argument(
        "--keep-threads", "-k",
        action="store_true",
        help="Keep threads after processing (default: delete threads)"
    )
    
    args = parser.parse_args()
    process_document(args.document_path, args.format, not args.keep_threads)

if __name__ == "__main__":
    # python process_document.py data/cv_doc_1.pdf
    main()